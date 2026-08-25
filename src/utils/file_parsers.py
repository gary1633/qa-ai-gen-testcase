import os
import re
from pathlib import Path
from typing import Tuple, List, Dict, Any, Optional
import docx
import pypdf
from src.integrations.jira_connector import JiraConnector, extract_jira_key

def clean_jira_key_from_title(title: str) -> str:
    """Loại bỏ mã Jira ticket (như VWCBT-3230) nếu bị dính vào trong câu tiêu đề testcase / scenario."""
    if not title:
        return ""
    # 1. Bỏ tiền tố mã Jira ở đầu câu: "[VWCBT-3230] " hoặc "VWCBT-3230: "
    cleaned = re.sub(r'^\s*\[?[A-Z0-9]+-\d+\]?[\s:-]+', '', title, flags=re.IGNORECASE)
    # 2. Bỏ mã Jira lọt vào giữa câu: "giao dịch VWCBT-3230 thành công" -> "giao dịch thành công"
    cleaned = re.sub(r'\s+\[?[A-Z0-9]+-\d+\]?\s*', ' ', cleaned, flags=re.IGNORECASE)
    # 3. Chuẩn hóa khoảng trắng
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned


def parse_markdown_or_text(file_path: str) -> str:
    """Đọc nội dung file Markdown hoặc Text thông thường."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read().strip()


def parse_docx_file(file_path: str) -> str:
    """Đọc file DOCX (PRD / SRS), trích xuất headings, paragraphs và tables."""
    doc = docx.Document(file_path)
    content_lines = []

    for elem in doc.paragraphs:
        text = elem.text.strip()
        if not text:
            continue
        if elem.style and elem.style.name and elem.style.name.startswith("Heading"):
            content_lines.append(f"\n## {text}\n")
        else:
            content_lines.append(text)

    if doc.tables:
        content_lines.append("\n### [Bảng thông tin / Tables]")
        for i, table in enumerate(doc.tables):
            content_lines.append(f"\n* Bảng {i+1}:")
            for row in table.rows:
                row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                content_lines.append("| " + " | ".join(row_cells) + " |")

    return "\n".join(content_lines).strip()


def parse_pdf_file(file_path: str) -> str:
    """Đọc file PDF, trích xuất text từ tất cả các trang."""
    reader = pypdf.PdfReader(file_path)
    pages_text = []
    for page_idx, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages_text.append(f"--- Trang {page_idx + 1} ---\n{text.strip()}")
    return "\n\n".join(pages_text).strip()


def is_safe_local_file(path_str: str) -> bool:
    """Kiểm tra chuỗi có phải là đường dẫn file hợp lệ trên ổ đĩa không (tránh lỗi File name too long khi là raw text)."""
    if not path_str or not isinstance(path_str, str):
        return False
    if "\n" in path_str or len(path_str) > 255:
        return False
    try:
        p = Path(path_str)
        return p.exists() and p.is_file()
    except (OSError, ValueError):
        return False


def extract_input_content(file_path_or_text: str) -> Tuple[str, str]:
    """
    Nhận diện file path, Jira Ticket URL / Key hoặc raw text và trả về (raw_text_content, file_type).
    """
    if not file_path_or_text:
        return "", "raw_text"

    clean_input = file_path_or_text.strip()

    # 1. Kiểm tra nếu là Jira Ticket URL hoặc chứa mã Jira Key (kèm ghi chú bổ sung nếu có)
    jira_key = extract_jira_key(clean_input)
    if jira_key and (not is_safe_local_file(clean_input)):
        try:
            jira_connector = JiraConnector()
            j_data = jira_connector.fetch_issue(jira_key)
            # Lấy phần text người dùng nhập thêm sau mã Jira (nếu có)
            extra_user_notes = re.sub(r'https?://[^\s]+|' + re.escape(jira_key), '', clean_input, flags=re.IGNORECASE).strip()
            extra_user_notes = re.sub(r'^[\s,;:-]+', '', extra_user_notes).strip()
            
            if extra_user_notes and len(extra_user_notes) >= 3:
                combined = (
                    f"{j_data['formatted_requirement']}\n\n"
                    f"---\n\n"
                    f"### THÔNG TIN BỔ SUNG / LÀM RÕ TỪ USER (USER CLARIFICATIONS & OVERRIDES):\n"
                    f"{extra_user_notes}"
                )
                return combined, "jira_with_clarifications"
            return j_data["formatted_requirement"], "jira"
        except Exception as e:
            # Nếu không kéo được Jira (chưa cấu hình token hoặc offline), fallback về raw text
            pass
    if is_safe_local_file(clean_input):
        path_obj = Path(clean_input)
        suffix = path_obj.suffix.lower()
        if suffix in [".md", ".markdown"]:
            return parse_markdown_or_text(clean_input), "md"
        elif suffix in [".txt", ".json", ".yaml", ".yml"]:
            return parse_markdown_or_text(clean_input), suffix.lstrip(".")
        elif suffix in [".docx"]:
            return parse_docx_file(clean_input), "docx"
        elif suffix in [".pdf"]:
            return parse_pdf_file(clean_input), "pdf"
        else:
            return parse_markdown_or_text(clean_input), "txt"
    else:
        # 3. User truyền thẳng raw text / user story string
        return clean_input, "raw_text"


def merge_multiple_sources(sources: List[str]) -> Tuple[str, str, Dict[str, Any]]:
    """
    Kết hợp nhiều nguồn tài liệu (Jira Tickets, File Word, PDF, OpenAPI, Text notes)
    thành một tài liệu phân tích nghiệp vụ tổng hợp duy nhất.
    """
    if not sources:
        return "", "empty", {}

    if len(sources) == 1:
        src_clean = sources[0].strip()
        jira_key = extract_jira_key(src_clean)
        if jira_key and (not is_safe_local_file(src_clean)) and ("/" in src_clean or len(src_clean.split()) <= 3):
            try:
                jira_connector = JiraConnector()
                j_data = jira_connector.fetch_issue(src_clean)
                metadata = {
                    "source_names": [f"Jira: {j_data['key']}"],
                    "jira_links": [j_data["jira_url"]],
                    "jira_keys": [j_data["key"]]
                }
                return j_data["formatted_requirement"], "jira", metadata
            except Exception as e:
                pass

        raw_text, f_type = extract_input_content(src_clean)
        metadata = {
            "source_names": [Path(src_clean).name if is_safe_local_file(src_clean) else "Direct_Input"],
            "jira_links": [],
            "jira_keys": [jira_key] if jira_key else []
        }
        return raw_text, f_type, metadata
    merged_parts = [
        "# TỔNG HỢP TÀI LIỆU YÊU CẦU ĐA NGUỒN (MULTI-DOCUMENT SPECIFICATION)",
        "> Hệ thống đã kết hợp các tài liệu nghiệp vụ, API specs và Jira tickets sau đây vào một luồng kiểm thử thống nhất:\n"
    ]

    metadata = {
        "source_names": [],
        "jira_links": [],
        "jira_keys": []
    }

    jira_connector = JiraConnector()

    for idx, src in enumerate(sources, 1):
        src_clean = src.strip()
        
        # 1. Kiểm tra nếu là Jira Ticket / URL
        jira_key = extract_jira_key(src_clean)
        if jira_key and (not is_safe_local_file(src_clean)) and ("/" in src_clean or len(src_clean.split()) == 1):
            try:
                j_data = jira_connector.fetch_issue(src_clean)
                merged_parts.append(f"## [Tài liệu {idx} - Jira Ticket]: {j_data['key']} - {j_data['summary']}")
                merged_parts.append(j_data["formatted_requirement"])
                metadata["jira_links"].append(j_data["jira_url"])
                metadata["jira_keys"].append(j_data["key"])
                metadata["source_names"].append(f"Jira: {j_data['key']}")
                continue
            except Exception as e:
                # Nếu lỗi kéo Jira, coi như raw text
                pass

        # 2. Kiểm tra nếu là Local File (.docx, .pdf, .md, .json, .yaml)
        if is_safe_local_file(src_clean):
            content, f_type = extract_input_content(src_clean)
            filename = Path(src_clean).name
            merged_parts.append(f"## [Tài liệu {idx} - File {f_type.upper()}]: {filename}")
            merged_parts.append(content)
            metadata["source_names"].append(filename)
        else:
            # 3. Raw text hoặc ghi chú / làm rõ bổ sung từ User
            merged_parts.append(f"## [Thông tin Bổ sung / Làm rõ từ User (User Clarifications & Overrides)]:")
            merged_parts.append(src_clean)
            metadata["source_names"].append(f"User_Clarification_{idx}")
    final_merged_text = "\n\n---\n\n".join(merged_parts)
    return final_merged_text, "multi_document", metadata
